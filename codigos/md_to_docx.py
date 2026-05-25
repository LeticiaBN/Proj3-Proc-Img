"""
Converte relatorio.md para relatorio.docx preservando:
  - headings (# ## ###)
  - tabelas markdown
  - negrito (**texto**)
  - blocos de codigo (``` ... ```)
  - listas (- item)
  - paragrafos normais
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


RAIZ = Path("/Users/leticia.neves/Desktop/6sem/2027/leo/entrega3")
SRC = RAIZ / "relatorio.md"
DST = RAIZ / "relatorio.docx"


def set_cell_bg(cell, hex_color="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_bold_italic(para, text):
    """Adiciona texto com suporte a **negrito** inline."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def parse_table(doc, lines):
    """Processa bloco de tabela markdown."""
    rows = []
    for line in lines:
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue  # linha separadora
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= n_cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            # remove **bold** markers for table cells
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
            run = para.add_run(clean)
            if ri == 0:
                run.bold = True
                set_cell_bg(cell, "BDD7EE")
    doc.add_paragraph()


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    i = 0
    while i < len(lines):
        line = lines[i]

        # --- bloco de codigo ---
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            para = doc.add_paragraph()
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
            para.paragraph_format.left_indent = Inches(0.4)
            i += 1
            continue

        # --- tabela ---
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            parse_table(doc, table_lines)
            continue

        # --- headings ---
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            content = re.sub(r"\*\*([^*]+)\*\*", r"\1", m.group(2))
            heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            doc.add_heading(content, level=level)
            i += 1
            continue

        # --- linha separadora ---
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        # --- lista com - ou * ---
        if re.match(r"^\s*[-*]\s+", line):
            content = re.sub(r"^\s*[-*]\s+", "", line)
            para = doc.add_paragraph(style="List Bullet")
            add_bold_italic(para, content)
            i += 1
            continue

        # --- linha vazia ---
        if line.strip() == "":
            i += 1
            continue

        # --- paragrafo normal ---
        para = doc.add_paragraph()
        add_bold_italic(para, line)
        i += 1

    doc.save(DST)
    print(f"Salvo: {DST}")


if __name__ == "__main__":
    main()
